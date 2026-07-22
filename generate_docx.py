import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_course_docx(filename):
    doc = Document()

    # Page setup - 2.54 cm margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette Definitions
    HEX_PRIMARY = "1A365D"    # Deep Navy
    HEX_SECONDARY = "2B6CB0"  # Slate Blue
    HEX_ACCENT = "C53030"     # Crimson / Emergency Red
    HEX_DARK = "2D3748"       # Dark Charcoal for text
    HEX_LIGHT_BG = "F7FAFC"   # Soft Gray Background
    HEX_WARNING_BG = "FFF5F5" # Soft Red Callout Background
    HEX_WARNING_BORDER = "E53E3E"
    HEX_NOTE_BG = "EBF8FF"    # Soft Blue Callout Background
    HEX_NOTE_BORDER = "3182CE"

    COLOR_PRIMARY = RGBColor(0x1A, 0x36, 0x5D)
    COLOR_SECONDARY = RGBColor(0x2B, 0x6C, 0xB0)
    COLOR_ACCENT = RGBColor(0xC5, 0x30, 0x30)
    COLOR_DARK = RGBColor(0x2D, 0x37, 0x48)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = COLOR_DARK
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper Functions for Formatting
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        return p

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.color.rgb = COLOR_DARK
        run_t = p.add_run(text)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.color.rgb = COLOR_DARK
        run_t = p.add_run(text)
        return p

    def set_cell_background(cell, hex_color):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_callout(text, title="AVISO IMPORTANTE", type="warning"):
        bg_color = HEX_WARNING_BG if type == "warning" else HEX_NOTE_BG
        border_color = HEX_WARNING_BORDER if type == "warning" else HEX_NOTE_BORDER
        title_color = COLOR_ACCENT if type == "warning" else COLOR_PRIMARY

        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        borders_xml = f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
        '''
        cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r_title = p.add_run(f"[{title}]\n")
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = title_color

        r_text = p.add_run(text)
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = COLOR_DARK
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def create_custom_table(doc, col_widths, headers, data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, title in enumerate(headers):
            hdr_cells[i].text = title
            set_cell_background(hdr_cells[i], HEX_PRIMARY)
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)

        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            bg_hex = HEX_LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = str(cell_value)
                set_cell_background(row_cells[col_idx], bg_hex)
                set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
                p = row_cells[col_idx].paragraphs[0]
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = COLOR_DARK

        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1A365D"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
        '''
        table._tbl.tblPr.append(parse_xml(borders_xml))
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return table

    # --- COVER / TITLE ---
    add_title("MANUAL DE INSTRUCCIÓN TEÓRICA Y OPERACIONAL")
    add_subtitle("Operación del Helicóptero MBB BO105 CBS4 para Servicios HEMS Offshore de Corto Alcance (< 8 km de Costa) — Aterrizaje en Helideck con Rotores Girando (Hot Loading)")

    # Intro Table Metadata
    meta_headers = ["Parámetro Operativo", "Especificación Mandatoria del Manual"]
    meta_data = [
        ["Aeronave y Configuración", "MBB BO105 CBS4 (MTOW 2.500 kg) — SIN Sistema de Flotación de Emergencia Fijo"],
        ["Condición de Vuelo", "Estrictamente VFR Diurno (Prohibición Absoluta de Vuelos Nocturnos e IMC)"],
        ["Perfil de Misión / Ocupantes", "HEMS Exclusivo (4 Ocupantes: Piloto PIC, Copiloto SIC, Médico Aeroevacuador, Paciente)"],
        ["Método de Evacuación", "Aterrizaje en Helideck del Buque y Embarque con Rotores en Movimiento (HOT LOADING). NO SE REALIZA IZAJE / WINCHING."],
        ["Límite de Exposición Overwater", "Inferior a 5 minutos acumulados en el perfil total de vuelo ida y vuelta"],
        ["Ubicación de Trabajos Marítimos", "DLV Seminole (Derrick Lay Vessel) a 7 km offshore frente a Punta Colorada, Río Negro"],
        ["Base Operativa / Destino Médico", "Base: Aeródromo Sierra Grande (SA21) / Destino HEMS: Aeropuerto Puerto Madryn (SAVY)"],
        ["Equipamiento Individual (PPE)", "Traje Antiexposición (Dry Suit), Chaleco Inflable Manual, PLB 406 MHz, Air Pocket Plus (EBS)"],
        ["Equipamiento Colectivo", "Balsa Salvavidas Inflable de 6 personas en cabina posterior a cargo del Médico Aeroevacuador"],
        ["Certificación de Tripulación", "Curso HUET (Helicopter Underwater Escape Training) Vigente para los 4 Ocupantes"]
    ]
    create_custom_table(doc, [2.5, 4.0], meta_headers, meta_data)

    add_callout(
        "RESTRICCIÓN OPERATIVA MANDATORIA: ESTA OPERACIÓN NO CONSIDERA IZAJE DE PACIENTE NI MANIOBRAS CON WINCHE/TORNO DE RESCATE.\n\n"
        "La totalidad de los embarques y desembarques del paciente y personal médico a bordo del buque DLV Seminole se realizarán ÚNICAMENTE mediante el posado de la aeronave sobre la cubierta del Helideck octogonal (22.2 m x 22.2 m), operando bajo el protocolo estricto de ROTORES EN MOVIMIENTO (HOT LOADING).\n\n"
        "La operación sin flotadores de emergencia fijos exige el cumplimiento 100% riguroso del límite de exposición overwater (< 5 min), del árbol de decisión Go/No-Go y de la portación del equipo individual de respiración subacuática Air Pocket Plus.",
        title="AVISO DE EXCLUSIÓN DE IZAJE Y PROCEDIMIENTO DE ATERRIZAJE (FOR TRAINING PURPOSES ONLY)",
        type="warning"
    )

    # --- MÓDULO 1 ---
    add_h1("MÓDULO 1: Marco Regulatorio y Definición Operativa Offshore Corto Alcance")
    add_h2("1.1 Marco Regulatorio ANAC / RAAC y Normativa Internacional")
    add_p("Las operaciones con helicópteros sobre superficies acuáticas se encuentran rigurosamente reguladas a nivel nacional e internacional para mitigar los riesgos inherentes a la pérdida de referencias visuales terrestres y la inmersión en caso de amerizaje forzoso (ditching). En la República Argentina, la Administración Nacional de Aviación Civil (ANAC) establece los marcos normativos en las Regulaciones Argentinas de Aviación Civil (RAAC):")
    add_bullet("Establece las reglas generales de vuelo, requerimientos de equipamiento mínimo y límites operativos para aeronaves de matrícula civil.", "RAAC Parte 91 (Reglas de Vuelo y Operación General): ")
    add_bullet("Fija las exigencias estrictas de entrenamiento de tripulaciones, mantenimiento preventivo y equipamiento de supervivencia para servicios de transporte y socorro aéreo médico.", "RAAC Parte 135 (Requerimientos Operacionales - Transporte Aéreo Comercial y HEMS): ")
    add_bullet("A nivel internacional, el Anexo 6 Parte III de la OACI y las especificaciones EASA SPA.HOFO gradúan los requerimientos de flotación fijos en función de la distancia a la costa y el tiempo de planeo disponible.", "Normativa OACI y EASA HOFO: ")

    add_h2("1.2 Definición Operativa de 'Corta Distancia' (< 8 km / 4.3 NM de Costa)")
    add_p("A los fines de este manual de instrucción, se define como 'Operación Offshore de Corto Alcance' a todo tramo de vuelo que se desarrolle sobre superficie acuática a una distancia igual o inferior a 8 kilómetros (4.3 millas náuticas) de una línea de costa continental apta para un aterrizaje forzoso seguro.")
    add_p("Esta definición es significativamente más restrictiva que el concepto general de 'Extended Overwater Operations'. La adopción de la cota de 8 km limita el radio geográfico a la franja de respuesta inmediata de los medios de búsqueda y rescate (SAR) costeros y Prefectura Naval Argentina.")

    add_h2("1.3 Exclusión Absoluta de Izaje (No-Hoist Operation) y Apoyo en Helideck")
    add_p("La configuración del helicóptero BO105 CBS4 para esta operación NO INCLUYE torno ni torno/winche de rescate (*No-Hoist Configuration*). La transferencia del paciente desde la embarcación hacia la aeronave se efectúa mediante toma de contacto de los patines de aterrizaje sobre la superficie despejada del Helideck del DLV Seminole.")
    add_p("Toda la tripulación (Piloto, Copiloto y Médico Aeroevacuador) debe estar habilitada en el protocolo de **Embarque/Desembarque con Rotores en Movimiento (*Hot Loading*)**, minimizando el tiempo de permanencia en cubierta y eliminando la necesidad de apagar los turbomotores en alta mar.")

    add_h2("1.4 Restricción Estricta a Condiciones VFR Diurnas Exclusivas")
    add_p("El presente protocolo operativo autoriza vuelos offshore ÚNICAMENTE en condiciones meteorológicas de vuelo visual (VFR) durante las horas de luz solar diurna. Queda ABSOLUTAMENTE PROHIBIDO bajo esta configuración de aeronave (sin flotadores de emergencia):")
    add_bullet("Vuelos nocturnos sobre agua (pérdida total de horizonte natural e imposibilidad de juzgar la altura sobre la superficie marina sin referencias terrestres).")
    add_bullet("Vuelos bajo Reglas de Vuelo por Instrumentos (IFR) o ingreso inadvertido en condiciones meteorológicas de instrumento (IIMC).")
    add_bullet("Operaciones de transporte comercial no-HEMS o maniobras de izaje suspendido con cable.", "Maniobras Prohibidas: ")

    add_h2("1.5 Justificación Técnica del Límite de Exposición Overwater (< 5 Minutos)")
    add_p("El pilar fundamental de la seguridad en esta operación radica en la drástica reducción del tiempo de exposición al riesgo sobre el mar. Al volar una aeronave bimotor a turbina (BO105 CBS4) sobre un tramo de agua de 7 km (3.8 NM), el tiempo de sobrevuelo puro sobre superficie líquida a velocidad de crucero (110 knots / 204 km/h) es de exactamente 2.1 minutos por pierna.")
    add_p("Considerando el tramo de ida desde la costa de Punta Colorada hacia el buque DLV Seminole (2.1 min) y el tramo de regreso desde el buque hacia la costa rumbo a Puerto Madryn (2.1 min), el tiempo acumulado total expuesto a la posibilidad de amerizaje forzoso es inferior a 5 minutos (4.2 min reales). Estadísticamente, la probabilidad de falla catastrófica simultánea de ambos motores o de la transmisión principal durante un intervalo de 4 minutos es de nivel 'Sumamente Improbable' (10^-7 a 10^-9 por hora de vuelo), lo que compensa la ausencia de flotadores estructurales.")

    # --- MÓDULO 2 ---
    add_h1("MÓDULO 2: Perfil Geográfico, Cartografía, Rutas y Planificación de Vuelo")
    add_h2("2.1 Cartografía de la Zona de Operación (Golfo San Matías y Golfo Nuevo)")
    add_p("La zona de operación comprende el litoral marítimo del sureste de la provincia de Río Negro y el noreste de la provincia del Chubut, sobre las aguas del Golfo San Matías y Golfo Nuevo:")
    add_bullet("Ubicado en S 41°36'00\" / W 65°21'00\". Pista de asfalto/tierra compactada apta para operaciones HEMS diurnas. Suministro de combustible Jet A-1 y punto de baseamiento de la aeronave.", "Base Principal (Aeródromo de Sierra Grande - SA21): ")
    add_bullet("Ubicada en S 41°42'00\" / W 65°01'00\". Punto de referencia terrestre continental y límite de transición entre el vuelo terrestre y el tramo overwater.", "Línea de Costa de Punta Colorada: ")
    add_bullet("Posicionado a 7 km (3.8 NM) mar adentro frente a Punta Colorada en S 41°43'12\" / W 64°56'24\". Buque de trabajo con Helideck octogonal Clase H2 de 22.2 m x 22.2 m.", "Embarcación Operativa (DLV Seminole): ")
    add_bullet("Ubicado en S 42°45'00\" / W 65°02'00\" (SAVY). Aeropuerto comercial asistido con infraestructura médica de alta complejidad en la ciudad de Puerto Madryn.", "Destino de Evacuación Médica (Puerto Madryn): ")

    add_h2("2.2 Desglose Fino de Tiempos y Distancias de Ruta")
    route_headers = ["Segmento de Vuelo", "Distancia (km / NM)", "Rumbo Mag", "Altitud Crucero", "Tiempo Vuelo (min)"]
    route_data = [
        ["Sierra Grande (SA21) -> Punta Colorada (Costa)", "32 km / 17.3 NM", "115°", "1.500 ft AGL", "9.4 min"],
        ["Punta Colorada (Costa) -> DLV Seminole (Offshore)", "7 km / 3.8 NM", "105°", "1.000 ft AGL", "2.1 min (OVERWATER)"],
        ["DLV Seminole -> Entrada Costa Punta Colorada", "7 km / 3.8 NM", "285°", "1.000 ft AGL", "2.1 min (OVERWATER)"],
        ["Punta Colorada (Costa) -> Puerto Madryn (SAVY)", "118 km / 63.7 NM", "185°", "1.500 ft AGL", "34.8 min"],
        ["TOTALES MISIÓN COMPLETA", "164 km / 88.6 NM", "—", "1.000 - 1.500 ft", "48.4 min (+ 5 min overwater)"]
    ]
    create_custom_table(doc, [2.2, 1.2, 0.9, 1.1, 1.1], route_headers, route_data)

    add_h2("2.3 Planificación de Combustible y Masa/Autonomía")
    add_p("El BO105 CBS4 cuenta con un sistema de combustible integrado por tanques principales en la estructura inferior y tanques de transferencia auxiliares con una capacidad útil total de 570 litros (456 kg de Jet A-1 a densidad 0.80 kg/L).")
    add_bullet("Consumo horario en crucero económico (110 KIAS): 220 L/h (176 kg/h).", "Consumo Específico: ")
    add_bullet("48.4 min de vuelo navegado + 10 min de embarque Hot Loading en helideck = 58.4 min totales (214 Litros / 171 kg).", "Combustible de Ruta: ")
    add_bullet("Exigencia de 30 minutos VFR diurno = 110 Litros (88 kg).", "Reserva Legal Mandatoria: ")
    add_bullet("214 L (Ruta) + 110 L (Reserva) = 324 Litros (259 kg).", "Carga Mínima al Despegue (Sierra Grande): ")
    add_bullet("Se despegará con 420 Litros (336 kg), otorgando un margen adicional de 25 minutos de contingencia médica.", "Carga Recomendada de Seguridad: ")

    # --- MÓDULO 3 ---
    add_h1("MÓDULO 3: Especificaciones Técnicas, Performance y Peso y Balance BO105 CBS4")
    add_h2("3.1 Características Generales y Sistema de Rotor Rígido (Rigid Rotor)")
    add_p("El MBB BO105 CBS4 es un helicóptero bimotor ligero caracterizado por su revolucionario sistema de rotor principal sin articulaciones de batimiento ni arrastre (Rotor Rígido System MBB), fabricado con una cabeza de titanio monolítica y 4 palas de plástico reforzado con fibra de vidrio (GFRP).")
    add_bullet("2x Turboshaft Allison 250-C20B con potencia máxima de despegue de 420 SHP cada uno.", "Planta Motriz: ")
    add_bullet("2.500 kg (5.511 lb). Versión CBS4 (cabina extendida 25 cm para configuración HEMS).", "Peso Máximo de Despegue (MTOW): ")
    add_bullet("La cabeza rígida proporciona una respuesta de control instantánea y un control de momento de encabritado/guio extremadamente potente, fundamental para contrarrestar la turbulencia y mantener la estabilidad durante el aterrizaje en helidecks marítimos.", "Respuesta Aerodinámica: ")

    add_h2("3.2 Degradación del Efecto Suelo (IGE vs OGE) en Aproximación Marítima")
    add_p("En operaciones sobre tierra firme, el colchón de aire comprimido generado por el rotor principal cerca del suelo (IGE - In Ground Effect) reduce la potencia requerida para el estacionario. Sin embargo, en sobrevuelo acuático previo a la toma en helideck:")
    add_bullet("La superficie marina no ofrece una almohadilla rígida; la presión del colchón de aire desplaza la masa de agua, generando olas concéntricas y spray.", "Pérdida de Compresión en Agua: ")
    add_bullet("Parte de la energía del rotor se disipa en la aceleración de la masa líquida descendente.", "Disipación de Energía: ")
    add_callout(
        "TODA PLANIFICACIÓN DE POTENCIA Y PESO PARA APROXIMACIONES, ESTACIONARIOS PREVIOS A LA TOMA O EMBARQUES SOBRE EL HELIDECK DEL DLV SEMINOLE DEBE REALIZARSE CALCULANDO VALORES OGE (OUT OF GROUND EFFECT).\n\n"
        "A 15°C ISA al nivel del mar, el techo OGE del BO105 CBS4 a MTOW (2.500 kg) requiere el 92% de Torque disponible. Se debe verificar el margen de potencia antes de iniciar el descenso final al helideck.",
        title="REGLA DE PERFORMANCE MANDATORIA EN APROXIMACIÓN AL HELIDECK",
        type="warning"
    )

    add_h2("3.3 Curva Altura-Velocidad (H-V Diagram) y Distancia de Planeo")
    add_p("La curva Altura-Velocidad del BO105 CBS4 define las zonas de sombra operativa (Avoid Area) donde un fallo de motor dificulta la autorrotación completa. En aproximación al helideck del buque:")
    add_bullet("La velocidad óptima de planeo y mínima tasa de descenso en autorrotación es de 65 KIAS.", "Velocidad de Autorrotación: ")
    add_bullet("La relación de planeo del BO105 CBS4 es de 4:1 (recorre 4 unidades horizontales por cada unidad de altura descendida).", "Relación de Planeo: ")
    add_bullet("A la altitud de crucero overwater recomendada de 1.000 ft AGL (~300 m), el helicóptero puede recorrer en autorrotación pura una distancia horizontal máxima de 1.2 km (0.65 NM).", "Cálculo de Distancia a Costa: ")
    add_bullet("Dado que el DLV Seminole está posicionado a 7 km (3.8 NM) de la costa, en el punto medio del tramo marítimo (3.5 km de la costa) el planeo en autorrotación NO ALCANZA LA TIERRA FIRME. El amerizaje forzoso (ditching) debe ser asumido como una maniobra planificada de alta probabilidad en caso de falla doble.", "Conclusión Crítica de Seguridad: ")

    # --- MÓDULO 4 ---
    add_h1("MÓDULO 4: Equipamiento de Supervivencia Individual (PPE) y Colectivo")
    add_h2("4.1 Traje Antiexposición (Dry Suit)")
    add_p("Las aguas del Golfo San Matías registran temperaturas medias anuales de entre 10°C y 14°C. La inmersión en agua a 10°C sin protección térmica desencadena la hipotermia severa y el choque térmico por frío en menos de 15 minutos, anulando la capacidad motriz en 5 minutos.")
    add_bullet("Cada ocupante debe portar un traje seco estanco, respirable, certificado bajo estándar ETTSO/TSO, equipado con sellos de látex o neopreno en cuello y muñecas.", "Requerimiento Técnico: ")
    add_bullet("Mantiene la temperatura corporal central por encima de 35°C durante un mínimo de 4 horas en agua a 10°C.", "Aislamiento Térmico: ")

    add_h2("4.2 Chaleco Salvavidas Inflable MANUAL")
    add_p("El chaleco salvavidas de aviación offshore debe cumplir con la especificación de doble cámara de inflado por cartuchos de CO2.")
    add_callout(
        "PROHIBICIÓN STRICTA DE CHALECOS CON SISTEMA DE INFLADO AUTOMÁTICO AL CONTACTO CON AGUA.\n\n"
        "Si el helicóptero acuatiza y se invierte (capsize) inundando la cabina, un chaleco de inflado automático se activará dentro del habitáculo. La flotabilidad positiva resultante atrapará a la persona contra el techo sumergido, impidiendo físicamente el egreso por las ventanas o puertas. El chaleco debe ser inflado ÚNICAMENTE LUEGO DE SALIR TOTALMENTE DE LA CABINA.",
        title="ADVERTENCIA VITAL DE SEGURIDAD OPERACIONAL",
        type="warning"
    )

    add_h2("4.3 Sistema de Respiración de Emergencia Subacuática (Air Pocket Plus / EBS)")
    add_p("El Air Pocket Plus es un sistema de rebreather de emergencia compacto que reutiliza el aire de la última espiración del ocupante en un pulmón de compensación flexible.")
    add_bullet("Proporciona entre 45 y 60 segundos de respiración subacuática autónoma en aguas frías.", "Autonomía de Respiración: ")
    add_bullet("Consta de boquilla de silicona con válvula unidireccional, pinza nasal integrada e inflador de purga rápida.", "Componentes: ")
    add_bullet("Elimina el reflejo de aspiración involuntaria de agua helada provocado por el choque térmico y la desorientación espacial.", "Función Táctica: ")

    add_h2("4.4 Radiobaliza de Localización Personal (PLB 406 MHz)")
    add_p("Fijada al arnés del chaleco salvavidas individual. Al alcanzar la superficie marina, se despliega la antena y se activa manualmente. Emite una señal de alerta satelital en 406 MHz integrada al sistema COSPAS-SARSAT con coordenadas GPS, y una señal de haming cercano en 121.5 MHz para las embarcaciones y aeronaves SAR de Prefectura Naval Argentina.")

    add_h2("4.5 Balsa Salvavidas Inflable de 6 Personas (6 Pax) en Cabina")
    add_p("Debido a la ausencia de flotadores estructurales en los patines, la aeronave lleva a bordo una balsa salvavidas inflable portátil homologada para 6 personas, alojada en el compartimento posterior de la cabina.")
    add_bullet("Ubicada inmediatamente detrás del asiento del médico aeroevacuador, asegurada con un arnés de desenganche rápido de un solo toque.", "Ubicación Física: ")
    add_bullet("Línea de amarre estructural (Painter Line) de 10 metros, que debe permanecer mosquetonada a un punto fijo de la cabina antes del lanzamiento para evitar que la balsa derive por el viento.", "Sistema de Amarre: ")
    add_bullet("El Médico Aeroevacuador es el responsable primario de sujetar el cabo de amarre, extraer la balsa al exterior tras el egreso en amerizaje y accionar el botellón de disparo de nitrógeno/CO2.", "Responsable de Despliegue: ")

    # --- MÓDULO 5 ---
    add_h1("MÓDULO 5: Operaciones en Cubierta del Buque DLV Seminole y Hot Loading")
    add_h2("5.1 Ficha Técnica y Especificaciones del Helideck del DLV Seminole")
    seminole_headers = ["Parámetro del Buque / Helideck", "Especificación Oficial del Operador (Micoperi)"]
    seminole_data = [
        ["Nombre de la Embarcación", "DLV Seminole (Derrick Lay Vessel - Buque de Tendido Submarino e Izaje)"],
        ["Dimensiones del Buque", "Eslora Total: 135.81 m / Manga: 30.5 m / Puntal: 9.0 m / Tonelaje: 13.871 GRT"],
        ["Sociedad de Clasificación", "RINA (Notaciones: HELIDECK, Pipe Laying Unit, Lifting Unit)"],
        ["Sistema de Posicionamiento", "Fondeo mediante 10 Líneas de Anclaje Electro-Hidráulicas (Buque NO DP)"],
        ["Forma y Dimensiones Helideck", "Octogonal de 22.2 metros × 22.2 metros (Área Despejada Amplia)"],
        ["Clasificación del Helideck", "Clase H2 (Apto para Helicópteros Medios y Pesados)"],
        ["Capacidad Máxima de Carga (MTOW)", "9.3 Toneladas (MTOW del BO105 es 2.5 t — Margen > 3.7x)"],
        ["Operación HEMS Autoraizada", "Toma de contacto en Helideck y Embarque Hot Loading (ROTORES GIRANDO). SIN IZAJE."]
    ]
    create_custom_table(doc, [2.5, 4.0], seminole_headers, seminole_data)

    add_h2("5.2 Implicancias del Sistema de Fondeo por 10 Anclas (Buque NO DP)")
    add_p("Al ser un buque de fondeo fijo y no contar con Posicionamiento Dinámico (DP), el DLV Seminole NO PUEDE virar su proa en forma rápida para ajustar el viento relativo a solicitud del piloto. El cambio de orientación requiere maniobrar el tendido de anclas (proceso de varias horas).")
    add_bullet("El Piloto al Mando debe solicitar por VHF el rumbo actual de la proa del buque y la lectura del anemómetro antes de iniciar la aproximación.", "Coordinación Previa: ")
    add_bullet("Si la orientación del buque genera viento de cola o viento cruzado fuerte sobre el helideck con turbulencia de la grúa principal de tendido del buque, el piloto evaluará el ángulo de aproximación en el sector despejado o suspenderá la maniobra de descenso.", "Evaluación de Viento Relativo: ")

    add_h2("5.3 Comunicaciones, Estado de Cubierta (Green Deck) y Protocolo Hot Loading")
    add_p("El contacto radial previo debe realizarse en VHF Marítima (Canal 16 o Canal de Trabajo designado) con el HLO (Helicopter Landing Officer) a 10 NM del buque. Se confirmará la condición 'Green Deck' y se preparará el embarque del paciente en helideck con rotores girando:")
    add_bullet("Cubierta despejada de obstáculos, redes de amarre tensionadas, cañones de espuma contra incendio en escucha, equipo HLO listo para guiar la camilla y movimientos de Pitch, Roll y Heave dentro de límites admisibles.", "Green Deck (Verde): ")
    add_bullet("Prohibición de aproximación o toma de contacto. El helicóptero orbitará a 0.5 NM en espera o abortará el descenso.", "Red Deck (Rojo): ")
    add_bullet("Una vez apoyados los patines en el Helideck, los pilotos mantendrán colectivo bajo y mandos centrados. El HLO autorizará al médico y al equipo de camilla del buque a aproximar a la aeronave por el sector frontal de seguridad (10:00 a 02:00) para embarcar al paciente y fijarlo a la camilla de cabina.", "Protocolo Hot Loading: ")

    # --- MÓDULO 6 ---
    add_h1("MÓDULO 6: Meteorología Costera y Marítima del Golfo San Matías")
    add_h2("6.1 Vientos Predominantes y Turbulencia Costera")
    add_p("En la costa de Punta Colorada predominan los vientos patagónicos del sector OESTE / SUROESTE. Al soplar desde la tierra firme hacia el mar (offshore wind):")
    add_bullet("La costa continental actúa como barrera, generando un mar calmo cerca de la orilla pero un oleaje picado y vientos de 25-35 nudos a 7 km mar adentro.", "Efecto Sombra Terrestre: ")
    add_bullet("Al cruzar la línea de costa, la diferencia de temperatura entre la tierra calentada por el sol y el agua fría genera una corriente descendente (downdraft) brusca.", "Cizalladura Costera: ")

    add_h2("6.2 Niebla de Advección Marina")
    add_p("Ocurre cuando masas de aire cálido del norte se desplazan sobre las aguas frías de la corriente marina en el Golfo San Matías. Produce bancos de niebla espesa con techo cero y visibilidad vertical nula que pueden cubrir el DLV Seminole en menos de 10 minutos.")
    add_p("Si la visibilidad cae por debajo de 3.000 metros en la zona del buque, se activará la orden de NO-GO o el retorno inmediato a la costa de Punta Colorada.")

    add_h2("6.3 Ilusión de Agua Calma (Glassy Water Illusion)")
    add_p("En días sin viento, el mar refleja el cielo como un espejo perfecto. Esto elimina los patrones de sombra de las olas y priva al ojo humano de referencias de profundidad y distancia vertical, provocando desorientación espacial o impacto inadvertido contra el agua durante el descenso al helideck. Se mitiga mediante el monitoreo estricto cruzado del radioaltímetro por el Copiloto.")

    # --- MÓDULO 7 ---
    add_h1("MÓDULO 7: Capacitación y Roles Específicos del Médico Aeroevacuador")
    add_h2("7.1 Gestión de Recursos Médicos Aéreos (AMRM)")
    add_p("El Médico Aeroevacuador en misiones offshore no es un mero pasajero; es un miembro activo de la tripulación de seguridad con responsabilidades específicas en la prevención de accidentes, coordinación en helideck y evacuación de emergencia.")

    add_h2("7.2 Protocolo de Embarque de Paciente en Helideck (Hot Loading)")
    add_p("Durante las operaciones HEMS con el helicóptero apoyado en el Helideck y rotores girando:")
    add_bullet("Ingresar a la superficie del helideck y aproximar a la aeronave ÚNICAMENTE tras recibir la señal visual de pulgar arriba del Piloto o HLO de cubierta.", "Aproximación Segura: ")
    add_bullet("Caminar en posición encorvada dentro del sector delantero visible (entre las 10:00 y las 02:00 del helicóptero).", "Sector Permitido: ")
    add_bullet("PROHIBIDO en todo momento circular hacia la parte posterior de los patines o hacia la zona del rotor de cola.", "Zona Mortal del Rotor de Cola: ")
    add_bullet("El paciente en camilla debe ser trasladado por el equipo de cubierta e ingresado por la puerta posterior/lateral, fijándolo inmediatamente con arnés de 4 puntos y asegurándole su chaleco salvavidas de activación manual.", "Aseguramiento del Paciente: ")

    add_h2("7.3 Rol del Médico en la Secuencia de Ditching y Egreso Subacuático")
    add_p("Ante el llamado verbal del Piloto '¡DITCHING, DITCHING, PREPARAR AIR POCKET!':")
    add_bullet("1. Colocar la boquilla del Air Pocket Plus en la boca, ajustar la pinza nasal y verificar la fijación del arnés personal.", "Paso 1: ")
    add_bullet("2. Palpar la manija de desenganche rápido de la Balsa de 6 Pax ubicada a su espalda.", "Paso 2: ")
    add_bullet("3. Verificar que la línea de amarre (Painter Line) de la balsa esté mosquetonada al punto fijo de la estructura de cabina.", "Paso 3: ")
    add_bullet("4. Mantener una mano fija en el marco de la puerta posterior como referencia espacial durante el acuatizaje y vuelco.", "Paso 4: ")
    add_bullet("5. Tras la detención del movimiento violento, accionar la manija de apertura de emergencia de la puerta posterior.", "Paso 5: ")
    add_bullet("6. Extraer la Balsa de 6 pax al exterior sujetando el cabo de amarre.", "Paso 6: ")
    add_bullet("7. Inflar su chaleco salvavidas LUEGO de salir de la cabina y accionar la botella de inflado de la balsa en el agua.", "Paso 7: ")
    add_bullet("8. Asistir el abordaje de los Pilotos y del Paciente a la balsa salvavidas desplegada.", "Paso 8: ")

    # --- MÓDULO 8 ---
    add_h1("MÓDULO 8: Procedimientos de Emergencia y Ditching Controlado")
    add_h2("8.1 Dinámica del Amaraje Forzoso (Ditching) en BO105 CBS4")
    add_p("Al carecer de flotadores fijos en los patines, la secuencia aerodinámica e hidrodinámica del contacto con el agua se desarrolla de la siguiente manera:")
    add_bullet("Se ejecuta una autorrotación de precisión manteniendo 65 KIAS hasta la superficie del mar.", "Contacto Inicial: ")
    add_bullet("Se realiza una levantada final (flare) para extinguir la velocidad horizontal, tocando el agua en actitud nivelada con el colectivo completo.", "Levantada Final: ")
    add_bullet("El centro de gravedad del BO105 CBS4 se ubica en la parte superior del fuselaje (motores Allison y transmisión de titanio). Al detenerse los patines en el agua, la masa superior genera un momento volcador instantáneo. LA AERONAVE SE DARÁ VUELTA E INVERTIRÁ (CAPSIZE) EN MENOS DE 5 SECUNDOS.", "Vuelco Inminente (Capsize): ")
    add_bullet("El Piloto aplicará paso colectivo máximo inmediatamente tras el contacto para frenar mecánicamente las palas del rotor principal contra el agua y reducir el peligro de rotación violenta.", "Frenado de Palas: ")

    add_h2("8.2 Secuencia Operativa HUET de Egreso Subacuático Paso a Paso")
    add_bullet("Mantener la calma y no soltar el arnés antes de fijar la mano de referencia en la ventana o puerta.", "Fase 1 (Contacto): ")
    add_bullet("Insertar la boquilla del Air Pocket Plus e iniciar la respiración pausada por la boca.", "Fase 2 (Inmersión): ")
    add_bullet("Esperar a que el agua inunde el habitáculo y el movimiento de rotación cese por completo.", "Fase 3 (Capsize): ")
    add_bullet("Accionar la palanca de desacople de emergencia de la puerta/ventana con la mano libre.", "Fase 4 (Apertura): ")
    add_bullet("Liberar la hebilla del arnés de pecho y deslizarse hacia afuera utilizando la mano de referencia como guía.", "Fase 5 (Egreso): ")
    add_bullet("Ascender a la superficie, inflar el chaleco de activación manual y activar la radiobaliza PLB de 406 MHz.", "Fase 6 (Superficie): ")

    # --- MÓDULO 9 ---
    add_h1("MÓDULO 9: Sistema de Gestión de Seguridad (SMS) y Matriz de Riesgo OACI")
    add_h2("9.1 Lista de Chequeo Mandatoria Go / No-Go HEMS Offshore")
    chk_headers = ["Condición Operativa Evaluada", "Criterio de Aprobación Mandatorio", "Estado Go / No-Go"]
    chk_data = [
        ["Reglas de Vuelo y Horario", "Visibilidad VFR > 5.000 m / Techo > 1.500 ft / Luz Diurna Solar", "OBLIGATORIO GO"],
        ["Límite de Distancia a Costa", "Posición del Buque / Trayectoria <= 8 km (4.3 NM) de Costa", "OBLIGATORIO GO"],
        ["Tiempo Exposición Overwater", "Tiempo acumulado de sobrevuelo acuático <= 5 minutos totales", "OBLIGATORIO GO"],
        ["Operación en Helideck", "Aterrizaje en Helideck y Embarque Hot Loading. SIN IZAJE/WINCHING.", "OBLIGATORIO GO"],
        ["Certificación HUET", "Curso HUET con validez vigente para los 4 Ocupantes a bordo", "OBLIGATORIO GO"],
        ["Equipamiento Individual PPE", "Dry Suits + Chalecos Manuales + PLB 406 MHz + Air Pocket Plus", "OBLIGATORIO GO"],
        ["Equipamiento Colectivo", "Balsa Salvavidas 6 Pax en cabina posterior con Painter Line fija", "OBLIGATORIO GO"],
        ["Estado del Helideck Buque", "Confirmación Radio VHF de GREEN DECK por el HLO del Seminole", "OBLIGATORIO GO"]
    ]
    create_custom_table(doc, [2.5, 3.0, 1.0], chk_headers, chk_data)

    add_h2("9.2 Matriz de Tolerabilidad de Riesgos OACI (Doc 9859)")
    risk_headers = ["Escenario de Peligro", "Riesgo Inicial", "Barreras Mitigadoras Aplicadas", "Riesgo Residual final"]
    risk_data = [
        ["Falla Técnica Overwater (Motor/Transmisión)", "3A (Inaceptable)", "• Exposición < 5 min\n• Bimotor Turbine\n• Mant. Preventivo Riguroso", "1B (Aceptable)"],
        ["Ahogamiento / Hipotermia tras Ditching", "3A (Inaceptable)", "• HUET Vigente\n• Dry Suits Aislantes\n• Air Pocket Plus (EBS)\n• Chalecos Manuales", "2C (Tolerable)"],
        ["Pérdida del Paciente durante Evacuación", "3A (Inaceptable)", "• Balsa 6 pax a cargo de Médico\n• Chaleco adaptado a paciente\n• Arnés de camilla 4 pts", "2B (Tolerable*)"],
        ["Ingreso Inadvertido en IIMC / Niebla Marina", "4A (Inaceptable)", "• VFR Diurno Exclusivo\n• Límite Visibilidad 5 km\n• Radioaltímetro Activo", "1A (Aceptable)"],
        ["Accidente en Helideck durante Hot Loading", "3B (Tolerable*)", "• Helideck Octogonal 22.2m\n• Sector de aproximación 10:00-02:00\n• Señalización HLO y no-izaje", "1C (Aceptable)"]
    ]
    create_custom_table(doc, [1.8, 1.0, 2.7, 1.0], risk_headers, risk_data)

    # --- MÓDULO 10 ---
    add_h1("MÓDULO 10: Syllabus de Evaluación y Cuestionario Teórico Extenso")
    add_h2("10.1 Cuestionario de Evaluación Teórica para Pilotos (25 Preguntas)")
    
    pilots_q = [
        ("1. ¿Cuál es el método exclusivo de transferencia del paciente autorizado a bordo del buque DLV Seminole?", 
         "a) Izaje con torno de rescate a 50 ft AGL.\nb) Transferencia por canasta suspendida desde grúa del buque.\nc) Toma de contacto en Helideck y embarque con rotores en movimiento (Hot Loading).\nd) Acuatizaje al costado del buque.", "c", 
         "El protocolo no considera izaje de paciente; la operación se realiza mediante aterrizaje en helideck y Hot Loading."),
        
        ("2. ¿Cuál es el tiempo máximo de exposición acumulado sobre agua permitido para el perfil completo ida y vuelta?", 
         "a) Inferior a 5 minutos totales.\nb) 15 minutos.\nc) 30 minutos.\nd) Sin límite si el vuelo es diurno.", "a", 
         "El tiempo de sobrevuelo overwater de ida (2.1 min) y vuelta (2.1 min) suma 4.2 min, reduciendo drásticamente la probabilidad de falla."),
        
        ("3. ¿En qué condiciones meteorológicas y de visibilidad está autorizada esta operación HEMS?", 
         "a) VFR diurno y nocturno.\nb) IFR con radar meteorológico activo.\nc) VFR especial nocturno.\nd) Estrictamente VFR Diurno exclusivo.", "d", 
         "La falta de flotadores prohíbe operaciones nocturnas e IFR por pérdida de horizonte natural sobre agua."),
        
        ("4. ¿Cuál es la velocidad recomendada de mínima tasa de descenso en autorrotación para el BO105 CBS4?", 
         "a) 90 KIAS.\nb) 65 KIAS.\nc) 45 KIAS.\nd) 110 KIAS.", "b", 
         "La velocidad de autorrotación aprobada en el RFM para mínima tasa de descenso es 65 KIAS."),
        
        ("5. ¿Por qué la planificación de potencia en estacionario o helideck marítimo debe hacerse calculando valores OGE?", 
         "a) Porque el downwash desplaza la masa de agua disipando la presión del colchón IGE.\nb) Porque el BO105 no posee certificación IGE.\nc) Por exigencia del proveedor de combustible Jet A-1.\nd) Para volar a mayor velocidad de crucero.", "a", 
         "El downwash desplaza la masa de agua disipando la presión del efecto suelo, requiriendo potencia OGE."),
        
        ("6. A 1.000 ft AGL de altitud de crucero sobre agua, ¿qué distancia horizontal máxima puede recorrer el BO105 en autorrotación?", 
         "a) 7 km.\nb) 5 km.\nc) 10 km.\nd) 1.2 km (0.65 NM).", "d", 
         "La relación de planeo 4:1 a 1.000 ft AGL (300 m) permite recorrer aproximadamente 1.2 km horizontales."),
        
        ("7. ¿Por qué está estrictamente PROHIBIDO el uso de chalecos salvavidas auto-inflables al contacto con el agua?", 
         "a) Porque pesan demasiado para la carga útil HEMS.\nb) Porque dañan el tapizado de la cabina.\nc) Porque el inflado intra-cabina atrapa al ocupante contra el techo sumergido impidiendo el egreso.\nd) Porque no funcionan en aguas frías.", "c", 
         "El inflado intra-cabina genera flotabilidad positiva inmediata que inmoviliza al tripulante contra el techo invertido."),
        
        ("8. ¿Qué autonomía de respiración subacuática proporciona el sistema Air Pocket Plus (EBS)?", 
         "a) 10 minutos.\nb) 45 a 60 segundos.\nc) 5 minutos.\nd) 30 segundos.", "b", 
         "El rebreather acumula el aire de la última espiración otorgando entre 45 y 60 segundos vitales bajo el agua."),
        
        ("9. ¿Cuál es la forma y dimensiones del Helideck del buque DLV Seminole?", 
         "a) Octogonal de 22.2 m × 22.2 m.\nb) Cuadrado de 10 m × 10 m.\nc) Circular de 15 m de diámetro.\nd) Rectangular de 30 m × 10 m.", "a", 
         "Las especificaciones oficiales RINA confirman una cubierta octogonal de 22.2 x 22.2 metros."),
        
        ("10. ¿Cuál es la capacidad máxima de carga útil del Helideck del DLV Seminole?", 
         "a) 2.5 toneladas.\nb) 15.0 toneladas.\nc) 9.3 toneladas.\nd) 5.0 toneladas.", "c", 
         "La clasificación Clase H2 del Seminole aprueba pesos de hasta 9.3 toneladas (amplio margen para los 2.500 kg del BO105)."),
        
        ("11. ¿Por qué el buque DLV Seminole no puede cambiar rápidamente su proa para ofrecer el viento relativo ideal al helicóptero?", 
         "a) Porque no tiene capitán habilitado.\nb) Porque el buque está varado en la orilla.\nc) Por falta de potencia en sus motores de propulsión.\nd) Porque está fondeado mediante 10 líneas de anclaje y no posee Posicionamiento Dinámico (DP).", "d", 
         "El sistema de fondeo de 10 anclas impide maniobras de viraje rápido de proa."),
        
        ("12. ¿Qué ocurre con la actitud del fuselaje del BO105 CBS4 inmediatamente tras acuatizar sin flotadores de emergencia?", 
         "a) Flota perfectamente nivelado en la superficie.\nb) Se volcará e invertirá (capsize) en menos de 5 segundos por tener el centro de gravedad elevado.\nc) Se hunde verticalmente sin volcar.\nd) Planea sobre la superficie marina.", "b", 
         "El peso superior de los motores Allison y transmisión genera un momento volcador instantáneo en agua."),
        
        ("13. ¿Qué acción debe realizar el piloto con el mando de paso colectivo inmediatamente tras el contacto con el agua en ditching?", 
         "a) Aplicar paso colectivo máximo para frenar mecánicamente las palas del rotor contra el agua.\nb) Bajar todo el colectivo inmediatamente.\nc) Soltar los mandos de vuelo.\nd) Mantener el colectivo en posición media.", "a", 
         "El paso colectivo máximo utiliza la resistencia del agua para frenar las palas y reducir la rotación violenta."),
        
        ("14. ¿En qué frecuencia emite la señal de alerta satelital la radiobaliza de localización personal (PLB)?", 
         "a) 118.1 MHz.\nb) 243.0 MHz.\nc) 406.0 MHz (con homing en 121.5 MHz).\nd) 156.8 MHz.", "c", 
         "El sistema Cospas-Sarsat opera en la frecuencia de alerta satelital de 406 MHz."),
        
        ("15. ¿Qué fenómeno óptico peligroso se produce sobre agua completamente calma en días sin viento?", 
         "a) Refracción solar directa.\nb) Ceguera nocturna.\nc) Espejismo de desierto.\nd) Ilusión de Agua Calma (Glassy Water Illusion) que elimina la percepción de altura.", "d", 
         "El reflejo especular del mar sin olas elimina las referencias de profundidad visual."),
        
        ("16. ¿Cuál es el Peso Máximo al Despegue (MTOW) certificado para el helicóptero BO105 CBS4?", 
         "a) 2.500 kg (5.511 lb).\nb) 2.200 kg.\nc) 3.000 kg.\nd) 1.800 kg.", "a", 
         "El peso máximo de despegue certificado del BO105 CBS4 es de 2.500 kg."),
        
        ("17. En la secuencia HUET de egreso subacuático, ¿cuándo se debe liberar el arnés de seguridad del asiento?", 
         "a) Antes de tocar la superficie del agua.\nb) LUEGO de fijar la mano de referencia en la puerta/ventana y verificar la detención del movimiento.\nc) Inmediatamente al volcar la aeronave.\nd) Nunca se debe liberar.", "b", 
         "Soltar el arnés sin mano de referencia provoca la pérdida total de orientación dentro de la cabina sumergida."),
        
        ("18. ¿Qué consumo promedio horario de Jet A-1 registra la planta motriz del BO105 CBS4 en crucero económico?", 
         "a) 100 L/h.\nb) 350 L/h.\nc) 220 L/h (176 kg/h).\nd) 500 L/h.", "c", 
         "La planta motriz bimotor Allison 250-C20B consume en promedio 220 litros por hora."),
        
        ("19. ¿Cuál es la función principal de la línea de amarre (Painter Line) de la balsa salvavidas de 6 personas?", 
         "a) Evitar que la balsa sea arrastrada por el viento o corriente tras ser arrojada al agua.\nb) Sujetar a los tripulantes al chaleco.\nc) Inflar la balsa dentro de la cabina.\nd) Amarrar el patín del helicóptero.", "a", 
         "La painter line mantiene la balsa vinculada a la aeronave o al rescatista en el agua."),
        
        ("20. ¿Qué rango de temperatura media registra el agua del Golfo San Matías?", 
         "a) 22°C a 25°C.\nb) 0°C a 4°C.\nc) 18°C a 20°C.\nd) 10°C a 14°C.", "d", 
         "Las frías aguas patagónicas del Golfo San Matías promedian entre 10°C y 14°C exigiendo traje seco secos."),
        
        ("21. ¿Qué información debe intercambiar el piloto con el HLO por VHF 10 NM antes de arribar al buque?", 
         "a) Menú de comidas a bordo.\nb) Rumbo/velocidad del buque, anemómetro, movimiento de cubierta y confirmación Green Deck.\nc) Número de pasaportes del paciente.\nd) La marca del helicóptero únicamente.", "b", 
         "Es el protocolo estándar de coordinación offshore para confirmar condiciones de helideck."),
        
        ("22. En el árbol de decisión Go/No-Go, si la visibilidad en la zona del buque se reduce a 2.000 metros, ¿cuál es la conducta?", 
         "a) Volar más bajo para buscar la cubierta.\nb) Continuar el vuelo por instrumentos IFR.\nc) RECHAZAR O ABORTAR LA MISIÓN (NO-GO).\nd) Apagar los equipos de radio.", "c", 
         "Al ser operación estricta VFR diurna, visibilidades menores a 3.000-5.000 m cancelan el vuelo."),
        
        ("23. ¿Qué componente del rotor principal del BO105 le proporciona su rigidez y respuesta de control instantánea?", 
         "a) Cabeza de rotor monolítica de titanio sin articulaciones de batimiento ni arrastre.\nb) Cabeza de rotor de madera laminada.\nc) Servomandos mecánicos únicamente.\nd) Palas de aluminio macizo.", "a", 
         "El Rotor Rígido MBB utiliza una cabeza monolítica de titanio de altísima resistencia."),
        
        ("24. ¿En qué frecuencia transmite la señal de radiogoniometría (homing) en cercanía del PLB para el rescate marítimo?", 
         "a) 156.8 MHz.\nb) 406.0 MHz.\nc) 118.0 MHz.\nd) 121.5 MHz.", "d", 
         "La frecuencia 121.5 MHz se utiliza para la radiogoniometría (homing) en proximidad por aeronaves SAR."),
        
        ("25. ¿Cuál es el destino de evacuación médica final establecido en la planificación de ruta para el paciente HEMS?", 
         "a) Helipuerto de Sierra Grande.\nb) Aeropuerto de Puerto Madryn (SAVY / El Tehuelche).\nc) Sanatorio de Comodoro Rivadavia.\nd) Hospital de Bahía Blanca.", "b", 
         "El protocolo fije el traslado directo al Aeropuerto de Puerto Madryn por su infraestructura asistencial.")
    ]

    for q_text, options, correct_ans, rationale in pilots_q:
        add_h3(q_text)
        add_p(options)

    add_h2("10.2 Cuestionario de Evaluación Teórica y Práctica para el Médico Aeroevacuador (15 Preguntas)")
    
    med_q = [
        ("1. En la rutina de seguridad AMRM, ¿dónde se ubica la Balsa Salvavidas de 6 personas en la cabina del BO105?", 
         "a) Estibada en el compartimento posterior de cabina, a espaldas del médico.\nb) En el baúl portaequipajes exterior.\nc) Debajo del asiento del piloto al mando.\nd) Amarrada al patín izquierdo de la aeronave.", "a", 
         "La balsa debe ser físicamente accesible en menos de 3 segundos por el médico desde su puesto."),
        
        ("2. ¿Cuándo debe el Médico Aeroevacuador colocarse la boquilla del Air Pocket Plus durante una emergencia sobre agua?", 
         "a) Al llegar a la superficie del mar tras nadar fuera.\nb) De inmediato al escuchar la voz de mando '¡DITCHING!' del piloto, ANTES del impacto en el agua.\nc) Después de liberar el arnés del asiento.\nd) No debe colocársela bajo ninguna circunstancia.", "b", 
         "El rebreather debe estar insertado en la boca antes de la inmersión para evitar la aspiración involuntaria de agua."),
        
        ("3. ¿Por qué el médico DEBE verificar que la Painter Line de la balsa esté mosquetonada antes del despegue?", 
         "a) Para que la balsa no haga ruido con la vibración.\nb) Para inmovilizar al paciente en la camilla.\nc) Para evitar que la balsa sea arrastrada por el viento o corriente tras ser arrojada al agua.\nd) No es necesario verificar la painter line.", "c", 
         "La painter line mantiene la balsa unida a la estructura o al médico evitando su pérdida en el mar."),
        
        ("4. Durante el embarque del paciente en helideck con rotores girando (Hot Loading), ¿por qué sector se debe aproximar a la aeronave?", 
         "a) Por la parte posterior cerca del rotor de cola.\nb) Por debajo de los patines de aterrizaje.\nc) Por cualquier sector sin mirar al piloto.\nd) Por el sector delantero visible entre las 10:00 y las 02:00 con indicación del piloto.", "d", 
         "El sector frontal garantiza que el piloto mantenga contacto visual constante con el personal médico."),
        
        ("5. ¿Cuándo debe el médico inflar su chaleco salvavidas individual durante la secuencia de egreso subacuático?", 
         "a) ÚNICAMENTE LUEGO DE SALIR COMPLETAMENTE DE LA CABINA AL EXTERIOR.\nb) Dentro de la cabina inundada antes de abrir la puerta.\nc) Mientras maniobra la apertura de la ventana de emergencia.\nd) Antes del impacto contra el agua.", "a", 
         "El inflado intra-cabina atrapa mecánicamente al médico contra el techo invertido impidiendo la salida."),
        
        ("6. ¿Qué función cumple la pinza nasal del sistema Air Pocket Plus (EBS)?", 
         "a) Mejorar la audición subacuática.\nb) Bloquear la vía aérea nasal para forzar la respiración exclusiva por la boquilla bajo el agua.\nc) Sujetar la máscara de oxígeno médico.\nd) Ninguna función táctica.", "b", 
         "Bloquea la vía aérea nasal para obligar la respiración exclusiva por la boquilla de silicona bajo el agua."),
        
        ("7. ¿Cuál es la primera acción del médico tras alcanzar la superficie marina luego del egreso subacuático?", 
         "a) Nadar de inmediato hacia la costa de Punta Colorada.\nb) Quitarse el traje seco antiexposición.\nc) Inflar el chaleco de activación manual, ubicar la painter line y accionar el inflado de la balsa de 6 pax.\nd) Gritar pidiendo ayuda sin inflar el chaleco.", "c", 
         "Garantiza su propia flotabilidad y activa la plataforma de supervivencia colectiva para el resto de los ocupantes."),
        
        ("8. ¿Por qué el traje seco (Dry Suit) es de uso obligatorio para el médico en vuelos sobre el Golfo San Matías?", 
         "a) Por estética del uniforme médico aéreo.\nb) Para protegerse de la radiación solar.\nc) Por exigencia del hospital de Puerto Madryn.\nd) Porque previene el choque térmico y la hipotermia severa en aguas de 10°C a 14°C.", "d", 
         "El choque térmico en agua a 10°C invalida físicamente a una persona en menos de 5 minutos sin traje seco."),
        
        ("9. ¿Cómo debe ser acondicionado el paciente HEMS en la camilla para el vuelo marítimo?", 
         "a) Sujetado con arnés de seguridad de 4 puntos y portando chaleco salvavidas de activación manual.\nb) Cubierto con una manta sin arnés para mayor comodidad.\nc) Suelto sobre la camilla para facilitar maniobras médicas.\nd) Atado con cintas de tela improvisadas.", "a", 
         "Asegura la retención del paciente durante turbulencias o impacto y provee flotabilidad individual."),
        
        ("10. En la cabina invertida y sumergida, si el médico pierde la visibilidad por agua turbia, ¿cómo debe orientarse para salir?", 
         "a) Soltando el arnés y pataleando en cualquier dirección.\nb) Manteniendo una mano fija de referencia en la estructura/puerta antes de liberar el cinturón.\nc) Esperando a que se encienda una luz de emergencia.\nd) Nadando a ciegas por el habitáculo.", "b", 
         "La mano de referencia es el único punto de anclaje propiocepción para saber dónde está la salida."),
        
        ("11. ¿Qué capacidad de carga y personas soporta la balsa salvavidas inflable portátil a bordo del BO105?", 
         "a) Máximo 2 personas.\nb) 10 personas únicamente.\nc) 6 personas (soporta holgadamente a los 4 ocupantes con equipamiento completo).\nd) 1 persona.", "c", 
         "Su diseño homologado soporta holgadamente el peso de 6 adultos con equipamiento completo."),
        
        ("12. ¿Qué señal visual o auditiva autoriza al médico a ingresar a la superficie del Helideck en el DLV Seminole?", 
         "a) Ninguna, se ingresa corriendo apenas aterriza.\nb) Encendido de luces rojas en la superestructura.\nc) Sirena continua del buque.\nd) Señal de pulgar arriba o contacto positivo del Piloto o HLO de cubierta.", "d", 
         "Es la norma de seguridad estricta para evitar ingresar a una cubierta no autorizada o inestable."),
        
        ("13. ¿Qué radiobaliza de localización personal (PLB) debe llevar fijada el médico a su chaleco salvavidas?", 
         "a) Baliza PLB de 406 MHz / 121.5 MHz de activación en superficie.\nb) Radio comercial AM/FM portátil.\nc) Teléfono celular común.\nd) Linterna estanca únicamente.", "a", 
         "La PLB permite a los satélites Cospas-Sarsat ubicar al médico y rescatistas en el agua."),
        
        ("14. Si el helideck del DLV Seminole se declara en estatus RED DECK por movimiento del buque, ¿cuál es la conducta del médico?", 
         "a) Exigir el aterrizaje inmediato por emergencia médica.\nb) Aceptar la espera o maniobra de seguridad indicada por el Piloto al Mando sin interferir.\nc) Abrir la puerta posterior en vuelo.\nd) Saltar desde el patín hacia la cubierta.", "b", 
         "El estatus Red Deck prohíbe operaciones por riesgo de vuelco o desprendimiento sobre la cubierta."),
        
        ("15. ¿Cuál es el tiempo máximo exigido al médico para colocarse la boquilla y pinza nasal del Air Pocket Plus en simulacros en tierra?", 
         "a) 30 segundos.\nb) 1 minuto.\nc) Menos de 10 segundos con ojos cerrados.\nd) 5 segundos sin ajustar la pinza nasal.", "c", 
         "El estándar de respuesta automatizada exige colocar boquilla y pinza en menos de 10 segundos.")
    ]

    for q_text, options, correct_ans, rationale in med_q:
        add_h3(q_text)
        add_p(options)

    add_h2("10.3 Pauta de Evaluación Práctica y Estándares de Vuelo")
    add_p("La habilitación final de las tripulaciones de vuelo y médicas requiere la firma conforme de la Planilla de Evaluación Práctica en Tierra y Vuelo (3 horas de simulación de autorrotación overwater, aproximación y toma en Helideck del DLV Seminole con Hot Loading y vuelo de evacuación HEMS a Puerto Madryn SAVY).")

    # Save Document
    filename_clean = r"c:\Users\SERVER-MADERO\Desktop\Output\cursoadptacionbo105bolkow\bo105-course\Curso_HEMS_Offshore_BO105_CBS4_Examen_Sin_Respuestas.docx"
    doc.save(filename_clean)
    print(f"Documento sin respuestas generado exitosamente en: {filename_clean}")
    
    try:
        doc.save(filename)
        print(f"Documento generado exitosamente en: {filename}")
    except PermissionError:
        alt_filename = filename.replace(".docx", "_v3.docx")
        doc.save(alt_filename)
        print(f"El archivo original estaba abierto en Word. Documento generado exitosamente en: {alt_filename}")

if __name__ == "__main__":
    out_path = r"c:\Users\SERVER-MADERO\Desktop\Output\cursoadptacionbo105bolkow\bo105-course\Curso_HEMS_Offshore_BO105_CBS4.docx"
    create_course_docx(out_path)
